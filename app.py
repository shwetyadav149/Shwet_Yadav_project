# SmartHire.AI Resume Builder — Version 2.0
# by Shwet Yadav
# ✨ Now includes: Projects, Certifications, ATS Upload, LinkedIn URL Analysis, and Word Download

import os
import gradio as gr
import google.generativeai as genai
import requests
import datetime
from docx import Document
import tempfile
import re

# =====================================
# CONFIGURATION
# =====================================
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
LOG_URL = os.getenv("SHEET_URL")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("models/gemini-2.5-flash")


# =====================================
# GOOGLE SHEET LOGGER
# =====================================
import json
import requests
import os

def log_to_google(user, module, status="Success", error=""):
    """Send log data to Google Sheets through Apps Script"""
    try:
        payload = {
            "user": user,
            "module": module,
            "status": status,
            "error": error
        }

        log_url = os.getenv("SHEET_URL")
        if not log_url:
            print("⚠️ SHEET_URL environment variable is missing.")
            return

        headers = {"Content-Type": "application/json"}
        print("📤 Sending log:", payload)

        response = requests.post(log_url, data=json.dumps(payload), headers=headers, timeout=10)

        print("📥 Google Sheet Response:", response.text)
        return response.text
    except Exception as e:
        print("❌ Error sending log:", e)


# =====================================
# GEMINI HELPER
# =====================================
def gemini_generate(prompt):
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"⚠️ Error: {e}"


# =====================================
# RESUME BUILDER
# =====================================
def build_resume(name, education, skills, experience, achievements, projects, certifications, languages, other_details):
    try:
        prompt = f"""
        Create a professional MBA Resume for {name}.
        Include these sections clearly formatted:
        - Education: {education}
        - Skills: {skills}
        - Experience: {experience}
        - Achievements: {achievements}
        - Projects: {projects}
        - Certifications: {certifications}
        - Languages: {languages}
        - Other Details: {other_details}
        Format it neatly with bullet points and section headers.
        """

        resume_text = gemini_generate(prompt)

        # Create downloadable Word file
        doc = Document()
        doc.add_heading(f"{name} — Professional Resume", level=1)
        for line in resume_text.split("\n"):
            doc.add_paragraph(line)
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_path.name)

        log_to_google(name, "Resume Builder", "Success", "")
        return resume_text, temp_path.name
    except Exception as e:
        log_to_google(name, "Resume Builder", "Failed", str(e))
        return f"⚠️ Error: {e}", None


# =====================================
# ATS SCORE CHECKER
# =====================================
def ats_check(uploaded_file, job_description):
    try:
        if uploaded_file is None:
            return "⚠️ Please upload your resume file first."

        # Extract text from uploaded PDF or DOCX
        resume_text = ""
        if uploaded_file.name.endswith(".pdf"):
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(uploaded_file.name)
            for page in pdf_reader.pages:
                resume_text += page.extract_text()
        elif uploaded_file.name.endswith(".docx"):
            from docx import Document
            doc = Document(uploaded_file.name)
            for para in doc.paragraphs:
                resume_text += para.text + "\n"
        else:
            return "⚠️ Please upload a PDF or Word (.docx) file."

        prompt = f"""
        Analyze this resume against the following job description.
        Resume: {resume_text}
        Job Description: {job_description}
        Provide:
        - ATS Compatibility Score (out of 100)
        - Key Missing Keywords
        - Recommendations for Improvement
        """

        result = gemini_generate(prompt)
        log_to_google("User", "ATS Score", "Success", "")
        return result
    except Exception as e:
        log_to_google("User", "ATS Score", "Failed", str(e))
        return f"⚠️ Error: {e}"


# =====================================
# LINKEDIN OPTIMIZER
# =====================================
def linkedin_optimize(profile_url_or_text):
    try:
        # If it's a LinkedIn URL, extract the public summary
        if re.match(r'https?://(www\.)?linkedin\.com', profile_url_or_text.strip()):
            result = f"""
            🔗 Analyzing LinkedIn URL: {profile_url_or_text}
            Since direct scraping is not allowed, please copy-paste your LinkedIn 'About' section text here next time.
            However, here’s a generic optimization guide:
            - Start with a 2-line personal branding intro.
            - Add quantifiable achievements.
            - End with a value statement for employers.
            """
        else:
            prompt = f"""
            Rewrite this LinkedIn 'About' section to make it more impactful, MBA-focused, and recruiter-friendly.
            Keep it under 2600 characters.
            Text: {profile_url_or_text}
            """
            result = gemini_generate(prompt)

        log_to_google("User", "LinkedIn Optimizer", "Success", "")
        return result
    except Exception as e:
        log_to_google("User", "LinkedIn Optimizer", "Failed", str(e))
        return f"⚠️ Error: {e}"


# =====================================
# GRADIO INTERFACE (With Student Login)
# =====================================
with gr.Blocks(theme=gr.themes.Soft()) as app:
    gr.Markdown(
        """
        <div style='text-align:center'>
            <h1>💼 SmartHire.AI — Resume Builder, ATS & LinkedIn Optimizer</h1>
            <p><b>By Shwet Yadav</b> | AI-Powered MBA Career Toolkit</p>
            <hr>
        </div>
        """
    )

    # 🔐 Login Section
    with gr.Group():
        student_name = gr.Textbox(label="👤 Enter Your Name or Student ID", placeholder="e.g., Shwet_Yadav_123")
        gr.Markdown("*(Please enter your name before using the tools below.)*")

    # --- Resume Builder Tab ---
    with gr.Tab("📄 Resume Builder"):
        name = gr.Textbox(label="Full Name")
        education = gr.Textbox(label="Education Background")
        skills = gr.Textbox(label="Key Skills")
        experience = gr.Textbox(label="Work Experience")
        achievements = gr.Textbox(label="Achievements")
        projects = gr.Textbox(label="Projects (optional)")
        certifications = gr.Textbox(label="Certifications (optional)")
        languages = gr.Textbox(label="Languages Known (optional)")
        other_details = gr.Textbox(label="Other Details (optional)")
        resume_output = gr.Textbox(label="Generated Resume", lines=15)
        download_file = gr.File(label="Download Word Resume")
        btn_resume = gr.Button("✨ Generate Resume")

        def resume_with_log(name, education, skills, experience, achievements, projects, certifications, languages, other_details, student_name):
            text, file = build_resume(name, education, skills, experience, achievements, projects, certifications, languages, other_details)
            log_to_google(student_name, "Resume Builder", "Success")
            return text, file

        btn_resume.click(
            resume_with_log,
            inputs=[name, education, skills, experience, achievements, projects, certifications, languages, other_details, student_name],
            outputs=[resume_output, download_file],
        )

    # --- ATS Checker Tab ---
    with gr.Tab("📊 ATS Score Checker"):
        resume_upload = gr.File(label="Upload Your Resume (PDF/DOCX)")
        job_desc = gr.Textbox(label="Paste Job Description", lines=10)
        ats_output = gr.Textbox(label="ATS Analysis Result", lines=15)
        btn_ats = gr.Button("🔍 Check ATS Score")

        def ats_with_log(resume_upload, job_desc, student_name):
            result = ats_check(resume_upload, job_desc)
            log_to_google(student_name, "ATS Checker", "Success")
            return result

        btn_ats.click(ats_with_log, inputs=[resume_upload, job_desc, student_name], outputs=ats_output)

    # --- LinkedIn Optimizer Tab ---
    with gr.Tab("🔗 LinkedIn Optimizer"):
        profile_input = gr.Textbox(label="Paste LinkedIn URL or 'About' Section", lines=6)
        linkedin_output = gr.Textbox(label="Optimized LinkedIn Version", lines=10)
        btn_linkedin = gr.Button("🚀 Optimize LinkedIn Profile")

        def linkedin_with_log(profile_input, student_name):
            result = linkedin_optimize(profile_input)
            log_to_google(student_name, "LinkedIn Optimizer", "Success")
            return result

        btn_linkedin.click(linkedin_with_log, inputs=[profile_input, student_name], outputs=linkedin_output)

    gr.Markdown(
        """
        <div style='text-align:center; margin-top:20px;'>
            <hr>
            <p>🚀 Built with ❤️ by <b>Shwet Yadav</b> | SmartHire.AI 2025</p>
        </div>
        """
    )

log_to_google("Shwet_Test", "System Test", "Success", "")
app.launch()

